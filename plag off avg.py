import os
import docx
def dicta(a):
    dict={}
    for i in a:
        if dict.get(i)==None:
            dict[i]=1
        else:
            dict[i]=dict[i]+1
    return dict
def sp(a):
    l=[]
    lm=[]
    for i in a:
        lm=i.split()
        for j in lm:
            #if (j.find(',') != -1) :
            j=j.replace("-","")
            j=j.replace(' ','')
            j=j.replace("/n"," ")
            j=j.replace(',','')
            j=j.replace("'","")
            j=j.lower()
            l.append(j)
    return l
        


asd=[]

bp=input("Enter the path where all '.docx' file are stored : ")
#pp=int(input('Enter the plagarism % filter (0-99):'))
s=''
for i in bp:
    if i=="\\":
        s=s+'/'
    else:
        s=s+i
if s[-1]=='/':
    print(s)
else:
    s=s+'/'
    print(s)        
    
basepath=s
for entry in os.listdir(basepath):
    if os.path.isfile(os.path.join(basepath,entry)):
        asd.append(entry)
#print(asd)
pans=[]
pavg=[]
common=['the','at','there','some','my','of','be','use','her','than','and','this','an','would','first','a','have','each','make','water','to','from','which','like','been','in','or','she','him','call','is','one','do','into','who','you','had','how','time','oil','that','by','their','has','its','it','word','if','look','now','he','but','will','two','find','was','not','up','more','long','for','what','other','write','down','on','all','about','go','day','are','were','out','see','did','as','we','many','get','with','when','then','no','come','his','your','them','way','made','they','can','these','could','may','I','said','so','part']
if len(asd)>=2:
    for i in range(len(asd)):
        doc1=[]
        doc1d={}
        d1=docx.Document(basepath+asd[i])
        for p in d1.paragraphs:
            doc1.append(p.text)
        doc1=sp(doc1)
        doc1dog=dicta(doc1)
        doc1d=dicta(doc1)
        plg=0
        s=0
        dog=0
        lkghj=[]
        for k in doc1d:
            s=s+doc1d[k]
        for j in range(len(asd)):
            if i!=j:
                #print(asd[i],'&',asd[j],'+++++++++++++')
                doc2=[]
                doc2d={}
                d2=docx.Document(basepath+asd[j])
                for p in d2.paragraphs:
                    doc2.append(p.text)
                doc2=sp(doc2)
                doc2d=dicta(doc2)
                lkgh=[]
                if doc1==doc2:
                    #print('********')
                    #print('Same',asd[i],'&',asd[j])
                    pans.append([asd[i],asd[j],100,doc1])
                    break
                
                else:

                    d=0  #same words
                    for k in doc1dog:
                        if (k in doc2d) and (not (k in common)) :
                            #print('y')
                            if (k not in lkgh):
                                lkgh.append(k)
                            if doc2d[k]-doc1dog[k]<=0:
                                d=d+doc2d[k]
                                
                            else:
                                d=d+doc1dog[k]
                        #else:
                            #d=doc1d[k]
                    plg=d/s
                    pans.append([asd[i],asd[j],round(plg*100,2),lkgh])
                
                for k in doc1d:
                    if (k in doc2d) and (not (k in common) and doc1d[k]>0 ):
                        if (k not in lkghj):
                            lkghj.append(k)
                            #print('y')
                        if doc2d[k]-doc1d[k]<=0:
                            doc1d[k]=doc1d[k]-doc2d[k]
                        else:
                            doc1d[k]=doc1d[k]-doc1d[k]
        dog=0
        for wd in doc1dog:
            dog=dog+(doc1dog[wd]-doc1d[wd])
            #print(doc1dog,doc1d)
        pavg.append([asd[i],'all',round((dog/s)*100,2),lkghj])
              
print('Max % copied from files---------------------------')
pmaxplg=[]
for i in asd:
    mxpg=0
    mxj=0
    for j in range (len(pans)):
        if i == pans[j][0]:
            if mxpg <= pans[j][2]:
                mxpg=pans[j][2]
                mxj=j
    pmaxplg.append(pans[mxj])
    #print('In ',pans[mxj][0],' from ',pans[mxj][1],' ---> ',pans[mxj][2],'%')
#print('--------------------------------------------------')

for i in pmaxplg:
    #print(i[0])
    for j in pavg:
        
        if i[0]==j[0]:
            if i[2]>j[2]:
                print('In ',i[0],' from ',i[1],' ---> ',i[2],'%-----------------------------------------------')
                print(i[3])
            else:
                print('In ',j[0],' from ',j[1],' ---> ',j[2],'%-----------------------------------------------')
                print(j[3])
print('--------------------------------------------------')
                

    
   
            
